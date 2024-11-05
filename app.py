import streamlit as st  
import openai  
from docx import Document  
import re  
from azure.ai.formrecognizer import DocumentAnalysisClient  
from azure.core.credentials import AzureKeyCredential  
  
# Azure OpenAI configuration  
AZURE_OPENAI_API_KEY = "783973291a7c4a74a1120133309860c0"  
AZURE_OPENAI_ENDPOINT = "https://theswedes.openai.azure.com/"  
OPENAI_API_TYPE = "azure"  
OPENAI_API_VERSION = "2024-05-01-preview"  
AZURE_DEPLOYMENT_NAME = "GPT-4-Omni"  
  
# Configure OpenAI API  
openai.api_key = AZURE_OPENAI_API_KEY  
openai.api_base = AZURE_OPENAI_ENDPOINT  
openai.api_type = OPENAI_API_TYPE  
openai.api_version = OPENAI_API_VERSION  
  
# Azure Form Recognizer setup  
form_recognizer_endpoint = "https://patentocr.cognitiveservices.azure.com/"  
form_recognizer_api_key = "cd6b8996d93447be88d995729c924bcb"  
  
def extract_claim_features(doc, claim_number):  
    claim_pattern = re.compile(rf"Independent Claim {claim_number}", re.IGNORECASE)  
    claim_features = []  
    found_claim = False  
  
    for para in doc.paragraphs:  
        if found_claim:  
            if para.text.strip():  
                claim_features.append(para.text.strip())  
            else:  
                break  
        if claim_pattern.search(para.text):  
            found_claim = True  
    return claim_features  
  
def extract_cited_reference_features(doc, reference_name):  
    reference_pattern = re.compile(rf"Cited Reference.*?{reference_name}", re.IGNORECASE)  
    reference_features = []  
    found_reference = False  
  
    for para in doc.paragraphs:  
        if found_reference:  
            if para.text.strip():  
                reference_features.append(para.text.strip())  
            else:  
                break  
        if reference_pattern.search(para.text):  
            found_reference = True  
    return reference_features  
  
def extract_text_from_pdf(pdf_data):  
    try:  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=pdf_data  
        )  
        result = poller.result()  
  
        text_content = {}  
        for page in result.pages:  
            page_text = ""  
            for line in page.lines:  
                page_text += line.content + "\n"  
            text_content[page.page_number] = page_text  
  
        return text_content  
    except Exception as e:  
        st.error(f"An error occurred during text extraction: {str(e)}")  
        return None  
  
def get_insights_from_llm(claim_features, cited_features, pdf_insights):  
    # Extract line numbers from cited features  
    line_number_pattern = re.compile(r'lines (\d+-\d+)')  
    cited_line_numbers = []  
  
    for feature in cited_features:  
        matches = line_number_pattern.findall(feature)  
        cited_line_numbers.extend(matches)  
  
    # Prepare the messages for the chat model  
    messages = [  
        {"role": "system", "content": "You are an expert in patent analysis."},  
        {  
            "role": "user",  
            "content": (  
                f"Extract and provide the content from the following line numbers in the PDF insights.\n"  
                f"Line Numbers: {cited_line_numbers}\n\n"  
                f"PDF Insights:\n{pdf_insights}\n\n"  
                "Instructions:\n"  
                "For each line number range, extract the text from the PDF insights that corresponds to those lines. "  
                "Organize the extracted text by each cited feature, displaying each one separately and in a clear format. "  
                "Ensure the extraction is accurate and precise."  
            )  
        }  
    ]  
  
    response = openai.ChatCompletion.create(  
        engine=AZURE_DEPLOYMENT_NAME,  
        messages=messages,  
        max_tokens=1500,  # Adjust if more content is needed  
        temperature=0.5  
    )  
      
    return response.choices[0].message['content'].strip()  
  
def main():  
    st.title("Patent Analysis Tool")  
  
    if 'word_file_uploaded' not in st.session_state:  
        st.session_state.word_file_uploaded = False  
        st.session_state.claim_features = []  
        st.session_state.reference_features = []  
        st.session_state.claim_number = ""  
        st.session_state.reference_name = ""  
        st.session_state.pdf_insights = {}  
  
    # Step 1: Upload Word file  
    uploaded_file = st.file_uploader("Upload a Word file", type=["docx"])  
  
    if uploaded_file is not None and not st.session_state.word_file_uploaded:  
        doc = Document(uploaded_file)  
  
        if st.button("Extract Features"):  
            st.session_state.claim_features = extract_claim_features(doc, st.session_state.claim_number)  
            st.session_state.reference_features = extract_cited_reference_features(doc, st.session_state.reference_name)  
            st.session_state.word_file_uploaded = True  
  
    if st.session_state.word_file_uploaded:  
        st.subheader(f"Key Features of Independent Claim {st.session_state.claim_number}")  
        if st.session_state.claim_features:  
            st.write(st.session_state.claim_features)  
        else:  
            st.write("No features found for the specified claim.")  
  
        st.subheader(f"Key Features of Cited Reference ({st.session_state.reference_name})")  
        if st.session_state.reference_features:  
            st.write(st.session_state.reference_features)  
        else:  
            st.write("No features found for the specified reference.")  
  
        # Step 2: Ask for number of PDF files to upload  
        num_files = st.number_input("How many PDF files do you want to upload?", min_value=1, step=1)  
  
        # Ensure that pdf_files is a list of the correct size  
        if 'pdf_files' not in st.session_state or len(st.session_state.pdf_files) != num_files:  
            st.session_state.pdf_files = [None] * num_files  
  
        all_insights_complete = True  
        for i in range(num_files):  
            pdf_file = st.file_uploader(f"Upload PDF file {i+1}", type=["pdf"], key=f"pdf{i+1}")  
  
            if pdf_file is not None:  
                pdf_data = pdf_file.read()  
                insights = extract_text_from_pdf(pdf_data)  
  
                if insights:  
                    st.session_state.pdf_insights[f"Document {i+1}"] = insights  
                    with st.expander(f"Insights for Document {i+1}"):  
                        for page_number, page_text in insights.items():  
                            st.write(f"Page {page_number}:")  
                            st.text(page_text)  
                    st.success(f"Insights stored for Document {i+1}")  
                else:  
                    st.warning(f"No insights extracted from PDF file {i+1}")  
                    all_insights_complete = False  
            else:  
                all_insights_complete = False  
  
        # Step 3: Generate insights using LLM  
        if all_insights_complete and st.button("Generate Result"):  
            comparison_result = get_insights_from_llm(  
                st.session_state.claim_features,  
                st.session_state.reference_features,  
                st.session_state.pdf_insights  
            )  
            st.write("Comparison Result:")  
            st.write(comparison_result)  
  
if __name__ == "__main__":  
    main()  
